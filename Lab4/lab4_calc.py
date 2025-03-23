# Comment on accuracy
#   Any sort of consistent error

from sympy import *
import numpy as np
import math
import copy
import matplotlib.pyplot as plt

from docx import Document

# outputing results to a word document
doc = Document()

# print a array to a word table
def array_to_word_table(array, name, decimals=4):
    doc.add_paragraph(name)
    table = doc.add_table(rows=len(array), cols=1)

    for i, cell in enumerate(array):
        table.cell(i, 0).text = str(round(cell, decimals))

# print a 2d array to a word table
def array2d_to_word_table(array, name, decimals=4):
    doc.add_paragraph(name)
    table = doc.add_table(rows=len(array), cols=len(array[0]))

    for i, row in enumerate(array):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(round(cell, decimals))

# define equations
b_omega, b_phi, b_kappa, b_lambda, tx, ty, tz = symbols('b_omega, b_phi, b_kappa, b_lambda, tx, ty, tz')
mx, my, mz = symbols('mx, my, mz')

# Create rotation matrix elements
m11 = (cos(b_phi)*cos(b_kappa))
m12 = ((sin(b_omega)*sin(b_phi)*cos(b_kappa)) + (cos(b_omega)*sin(b_kappa)))
m13 = (-(cos(b_omega)*sin(b_phi)*cos(b_kappa)) + (sin(b_omega)*sin(b_kappa)))

m21 = (-cos(b_phi)*sin(b_kappa))
m22 = (-(sin(b_omega)*sin(b_phi)*sin(b_kappa)) + (cos(b_omega)*cos(b_kappa)))
m23 = ((cos(b_omega)*sin(b_phi)*sin(b_kappa)) + (sin(b_omega)*cos(b_kappa)))

m31 = (sin(b_phi))
m32 = (-sin(b_omega)*cos(b_phi))
m33 = (cos(b_omega)*cos(b_phi))

# Populate Matrix
rotation_matrix = Matrix([
    [m11, m12, m13],
    [m21, m22, m23],
    [m31, m32, m33],
])

class Absolute_Orientation:
    # absolute orientation equations
    O = b_lambda * rotation_matrix * Matrix([mx, my, mz]) + Matrix([tx, ty, tz])
    xO, yO, zO = O

    partial_derivatives = []
    for o in O:
        row = []
        row.append(diff(o, b_omega))
        row.append(diff(o, b_phi))
        row.append(diff(o, b_kappa))
        row.append(diff(o, b_lambda))
        row.append(diff(o, tx))
        row.append(diff(o, ty))
        row.append(diff(o, tz))
        partial_derivatives.append(row)

    def __init__(self, name, model, object, base, rel_ori):
        self.model = model
        self.object = object
        self.name = name

        #self.plot()

        array2d_to_word_table(self.model, f"{self.name} input model")
        array2d_to_word_table(self.object, f"{self.name} input object")

        # Getting parameters
        self.inital_approx(np.array(self.object[0]), np.array(self.object[1]), np.array(self.model[0]), np.array(self.model[1]))
        self.set_deg_parameters()
        array_to_word_table(self.deg_parameters, f"{self.name} Initial Parameters")
        self.estimate_parameters()
        self.set_deg_parameters()
        print(f"Final parameters {self.deg_parameters}")
        array_to_word_table(self.deg_parameters, f"{self.name} Parameters")

        # object coordinates and residuals
        self.transformed_points = self.get_ground_coordinates(self.model)
        self.residuals = self.get_residuals(self.transformed_points, self.object)
        array2d_to_word_table(self.transformed_points, f"{self.name} transformed model points")
        array2d_to_word_table(self.residuals, f"{self.name} residuals for transformed model points")

        # transforming perspective centers
        left_image_pc = self.get_ground([0, 0, 0])
        right_image_pc = self.get_ground(base)
        array2d_to_word_table([left_image_pc], f"{self.name} left perspective center")
        array2d_to_word_table([right_image_pc], f"{self.name} right perspective center")

        print(f"baseline distance: {np.linalg.norm(np.array(left_image_pc) - np.array(right_image_pc))}")

        extracted_angles = self.extract_angles(rel_ori, self.parameters)
        print(extracted_angles)
        array2d_to_word_table(extracted_angles, f"{self.name} extracted angles")

    def inital_approx(self, o1, o2, m1, m2):
        ao = atan((o2[0] - o1[0])/(o2[1] - o1[1]))
        am = atan((m2[0] - m1[0])/(m2[1] - m1[1]))

        kappao = ao - am
        
        # r3
        Mo = np.array([[np.cos(float(kappao)),  np.sin(float(kappao)), 0],
                    [-np.sin(float(kappao)), np.cos(float(kappao)), 0],
                    [0, 0, 1]])

        delta_xo = o2[0] - o1[0]
        delta_yo = o2[1] - o1[1]
        delta_zo = o2[2] - o1[2]

        delta_xm = m2[0] - m1[0]
        delta_ym = m2[1] - m1[1]
        delta_zm = m2[2] - m1[2]

        do = math.sqrt(delta_xo**2 + delta_yo**2 + delta_zo**2)
        dm = math.sqrt(delta_xm**2 + delta_ym**2 + delta_zm**2)

        lambdao = do / dm

        to = o1.T - (lambdao * (Mo @ o2.T))

        self.parameters = np.array([0.0, 0.0, kappao, lambdao, to[0], to[1], to[2]]) # omega, phi, kappa, lambda, tx, ty, tz
        print(f"Initial parameters: {self.parameters}")

    # points [x, y, z]
    def design_matrix(self):
        A = []
        misclosure_vector = []
        for point_index in range(len(self.model)):
            for coord_index in range(3):
                A_row = []
                
                for p in self.partial_derivatives[coord_index]:
                    value = self.evaluate(p, self.model[point_index])
                    A_row.append(value)

                A.append(A_row)
            misclosure_vector += self.misclosure(self.model[point_index], self.object[point_index])
        return np.matrix(A), misclosure_vector

    # does the LSE to get parameters for absolute orientation
    def estimate_parameters(self):
        iterations = 0
        while(1):
            A, misclosure_vector = self.design_matrix()

            corrections = -(np.linalg.inv(A.T @ A) @ A.T @ np.array(misclosure_vector))
            self.parameters = self.parameters + np.array(corrections).ravel()

            rel_corrections = np.abs(corrections / (np.abs(self.parameters) + 1e-10))
            iterations += 1
            if np.max(rel_corrections) < 1e-3: # if converged
                print(f"Iterations: {iterations}")
                A, misclosure_vector = self.design_matrix()
                print(f"Residuals: {np.array(misclosure_vector)}")

                # correlation matrix
                cov = np.linalg.inv(A.T @ A)
                std_devs = np.sqrt(np.diag(cov))  # Standard deviations of parameters
                corr_matrix = cov / np.outer(std_devs, std_devs)
                print("Correlation Matrix of Estimated Parameters:\n", corr_matrix)
                array2d_to_word_table(corr_matrix.tolist(), f"{self.name} Correlation Matrix")

                # Compute hat matrix and redundancy numbers
                H = A @ np.linalg.inv(A.T @ A) @ A.T
                redundancy_numbers = 1 - np.diag(H)  # Redundancy numbers for each observation
                redundancy_table = redundancy_numbers.reshape(len(self.model), 3)

                print("Redundancy Numbers:\n", redundancy_numbers)
                array2d_to_word_table(redundancy_table.tolist(), f"{self.name} Redundancy Numbers", decimals=4)

                break

    def evaluate(self, eq, point):
        new_eq = eq.subs({
            b_omega: self.parameters[0],
            b_phi: self.parameters[1], 
            b_kappa: self.parameters[2], 
            b_lambda: self.parameters[3], 
            tx: self.parameters[4], 
            ty: self.parameters[5], 
            tz: self.parameters[6],
            mx: point[0],
            my: point[1],
            mz: point[2]
        })
        
        return float(N(new_eq, 50))

    def misclosure(self, model_coord, object_coord):
        w = [
            self.evaluate(self.xO, model_coord) - object_coord[0],
            self.evaluate(self.yO, model_coord) - object_coord[1],
            self.evaluate(self.zO, model_coord) - object_coord[2]
        ]
        return w

    def get_ground(self, coord):
        x = self.evaluate(self.xO, coord)
        y = self.evaluate(self.yO, coord)
        z = self.evaluate(self.zO, coord)
        return [x, y, z]

    def get_ground_coordinates(self, model_points):
        transformed_points = []
        for coord in model_points:
            transformed_coord = self.get_ground(coord)
            transformed_points.append(transformed_coord)

        return transformed_points

    def get_residuals(self, transformed_points, ground_points):
        residuals = []
        for i in range(len(transformed_points)):
            x = transformed_points[i][0] - ground_points[i][0]
            y = transformed_points[i][1] - ground_points[i][1]
            z = transformed_points[i][2] - ground_points[i][2]
            residuals.append([x, y, z])

        # compute MeanError and RMSE
        residuals = np.array(residuals)
        mean_error = np.mean(np.abs(residuals), axis=0)  # Compute mean error for each column
        rmse = np.sqrt(np.mean(np.square(residuals), axis=0))
        residuals = np.vstack([residuals, mean_error, rmse])
        
        return residuals

    def evaluate_rotation(self, parameters):
        new_eq = rotation_matrix.subs({
            b_omega: parameters[0],
            b_phi: parameters[1], 
            b_kappa: parameters[2], 
        })
        
        return N(new_eq, 50)

    def set_deg_parameters(self):
        self.deg_parameters = copy.copy(self.parameters)
        self.deg_parameters[0] *= 180 / math.pi
        self.deg_parameters[1] *= 180 / math.pi
        self.deg_parameters[2] *= 180 / math.pi

    def get_rad_parameters(self, p):
        dx = p[0] * math.pi / 180
        dy = p[1] * math.pi / 180
        dz = p[2] * math.pi / 180
        return [dx, dy, dz]

    def get_angle(self, M):
        print(M)
        omega = np.degrees(np.arctan2(-M[2, 1], M[2, 2]))
        phi = np.degrees(np.arcsin(M[2, 0]))
        kappa = np.degrees(np.arctan2(-M[1, 0], M[0, 0]))
        return [omega, phi, kappa]

    # extract angles from the relative parameters and the absolutae parmeters
    # the angles for going from image to object space
    def extract_angles(self, rel_p, abs_p):
        array2d_to_word_table([rel_p], f"{self.name} relative orientation parameters")
        M_i_m_left = self.evaluate_rotation(self.get_rad_parameters([0, 0, 0]))
        M_i_m_right = self.evaluate_rotation(self.get_rad_parameters(rel_p)) # relative orientation rotation parameters
        array2d_to_word_table(M_i_m_left.tolist(), f"{self.name} M_i_m_left")
        array2d_to_word_table(M_i_m_right.tolist(), f"{self.name} M_i_m_right")

        M_o_m = self.evaluate_rotation(abs_p)
        array2d_to_word_table(M_o_m.tolist(), f"{self.name} M_o_m")
        array2d_to_word_table(M_o_m.T.tolist(), f"{self.name} M_o_m.T")

        M_i_o_left = M_i_m_left * M_o_m.T
        M_i_o_right = M_i_m_right * M_o_m.T
        array2d_to_word_table(M_i_o_left.tolist(), f"{self.name} M_i_o_left")
        array2d_to_word_table(M_i_o_right.tolist(), f"{self.name} M_i_o_right")

        angles = []
        angles.append(self.get_angle(np.array(M_i_o_left).astype(np.float64)))
        angles.append(self.get_angle(np.array(M_i_o_right).astype(np.float64)))
        return angles

    def plot(self):
        points = np.array(self.model)

        plt.figure()

        plt.plot(points[:, 0], points[:, 1], 'ro')
        plt.title('GCPs in model space')
        plt.xlabel('x (mm)')
        plt.ylabel('y (mm)')
        plt.grid(True)
        plt.axis([-125, 125, -125, 125])

        plt.tight_layout()
        plt.show()

def main():

    # lab data
    print("\n\nLab Data\n")

    object_coords = np.array([
        [-399.28, -679.72, 1090.96], # 100
        [109.70, -642.35, 1086.43],  # 102
        [475.55, -538.18, 1090.50],  # 104
        [517.62, -194.43, 1090.65],  # 105
        [-466.39, -542.31, 1091.55], # 200
        [42.73, -412.19, 1090.82],   # 201
        [321.09, -667.45, 1083.49],  # 202
        [527.78, -375.72, 1092.00]   # 203
    ])
    '''
    all_model_points = np.array([
        [-9.5921, 96.2715, -158.3930],  # 100
        [-2.3846, -5.9159, -156.4915],  # 102
        [18.3668, -79.3166, -153.5411], # 104
        [87.3462, -88.0244, -153.0441], # 105
        [18.1668, 109.7020, -158.5132], # 200
        [43.8732, 7.3538, -155.7842],   # 201
        [-7.5412, -48.3543, -155.9639], # 202
        [50.8805, -89.9778, -152.9167]  # 203
    ])
    '''
    all_model_points = np.array([
        [-9.5975,	96.3215,	-153.4711], # 100
        [-2.3859,	-5.9182,	-151.6286], # 102
        [18.3766,	-79.3583,	-148.7698], # 104
        [87.3915,	-88.0704,	-148.2892], # 105
        [18.1762,	109.7568,	-153.5850], # 200
        [43.8954,	7.3577, 	-150.9432], # 201
        [-7.5452,	-48.3790,	-151.1187], # 202
        [50.9073,	-90.0254,	-148.1656] # 203
    ])

    array2d_to_word_table(object_coords, f"all object")
    array2d_to_word_table(all_model_points, f"all model")

    selected_indices = [1, 2, 3, 4] # what indices are control points

    # extracting control and check points
    control_points = all_model_points[selected_indices]
    check_points = np.delete(all_model_points, selected_indices, axis=0)
    control_object_coords = object_coords[selected_indices]
    check_object_points = np.delete(object_coords, selected_indices, axis=0)

    print(f"Control Points {control_points}")

    lab_abs = Absolute_Orientation("Lab", control_points, control_object_coords, [92, -1.4649, -1.2609], [-0.9724, 0.251, -1.7526])

    # transforming check points
    check_transformed = lab_abs.get_ground_coordinates(check_points)
    check_residuals = lab_abs.get_residuals(check_transformed, check_object_points)
    array2d_to_word_table(check_transformed, f"check points in object")
    array2d_to_word_table(check_residuals, f"check points residuals")
    print(f"Check Residuals:\n{check_residuals}")

    # transforming tie points
    tie_points = np.array([
        [-9.9688, 14.8164, -156.2243],    # T1
        [92.0658, -4.0001, -154.4931],    # T2
        [-10.5394, -102.5484, -155.0824], # T3
        [87.0929, -88.3482, -153.1153],   # T4
        [-9.4881, 96.2460, -158.3431],    # T5
        [84.9816, 102.8444, -157.5453]    # T6
    ])

    array2d_to_word_table(tie_points, f"all tie")

    tie_transformed = lab_abs.get_ground_coordinates(tie_points)
    array2d_to_word_table(tie_transformed, f"tie points in object")

    # test data
    print("\n\nTest Data\n")
    model_coordinates = [
        [108.9302,  92.5787, -155.7696], # 30
        [ 19.5304,  96.0258, -156.4878], # 40
        [ 71.8751,   4.9657, -154.1035], # 72
        [ -0.9473,  -7.4078, -154.8060], # 127
        [  9.6380, -96.5329, -158.0535], # 112
        [100.4898, -63.9177, -154.9389], # 50
    ]

    object_coordinates = [
        [7350.27,	4382.54,	276.42], # 30
        [6717.22,	4626.41,	280.05], # 40
        [6869.09,	3844.56,	283.11], # 72
        [6316.06,	3934.63,	283.03], # 127
        [6172.84,	3269.45,	248.10], # 112
        [6905.26,	3279.84,	266.47], # 50
    ]

    test_data = Absolute_Orientation("Test", model_coordinates, object_coordinates, [92, 5.0455, 2.1725], [0.4392, 1.508, 3.1575])

    # printing answers out to a word document
    doc.save("output.docx")

#print(test_data.get_ground([92, 5.0455, 2.1725]))
if __name__ == "__main__":
    main()