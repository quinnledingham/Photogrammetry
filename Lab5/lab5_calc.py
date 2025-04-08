from sympy import *
import numpy as np
import math
from docx import Document
import copy
import sys
import pandas as pd
import pickle
import matplotlib.pyplot as plt

sys.path.insert(1, './../Lab2')
sys.path.insert(1, './Lab2')

import lab2_calc

object_coords = np.array([
    [-399.28, -679.72, 1090.96], # 100
    [109.70, -642.35, 1086.43],  # 102
    [475.55, -538.18, 1090.50],  # 104
    [517.62, -194.43, 1090.65],  # 105
    [-466.39, -542.31, 1091.55], # 200
    [42.73, -412.19,  1090.82],  # 201
    [321.09, -667.45, 1083.49],  # 202
    [527.78, -375.72, 1092.00]   # 203
])

# outputing results to a word document
doc = Document()

images = lab2_calc.main()

#with open("test", "wb") as fp:
#    pickle.dump(images, fp)

#with open("./test", "rb") as fp:
#    images = pickle.load(fp)

images[0].rms *= np.sqrt(20)
images[1].rms *= np.sqrt(25)

# print a array to a word table
def array_to_word_table(array, name, decimals=4):
    doc.add_paragraph(name)
    table = doc.add_table(rows=len(array), cols=1)

    for i, cell in enumerate(array):
        table.cell(i, 0).text = str(round(cell, decimals))

# print a 2d array to a word table
def array2d_to_word_table(array, name, decimals=4):
    doc.add_paragraph(name)
    table = doc.add_table(rows=len(array), cols=len(array[0]))

    for i, row in enumerate(array):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(round(cell, decimals))

Xc, Yc, Zc, omega, phi, kappa = symbols('Xc, Yc, Zc, omega, phi, kappa')
Xi, Yi, Zi, a, ix, b, iy, tx, ty, c, Xp, Yp = symbols('Xi, Yi, Zi, a, ix, b, iy, tx, ty, c, Xp, Yp')
xij_bar, yij_bar = symbols('xij_bar, yij_bar')

# Create rotation matrix elements
m11 = (cos(phi)*cos(kappa))
m12 = ((sin(omega)*sin(phi)*cos(kappa)) + (cos(omega)*sin(kappa)))
m13 = (-(cos(omega)*sin(phi)*cos(kappa)) + (sin(omega)*sin(kappa)))

m21 = (-cos(phi)*sin(kappa))
m22 = (-(sin(omega)*sin(phi)*sin(kappa)) + (cos(omega)*cos(kappa)))
m23 = ((cos(omega)*sin(phi)*sin(kappa)) + (sin(omega)*cos(kappa)))

m31 = (sin(phi))
m32 = (-sin(omega)*cos(phi))
m33 = (cos(omega)*cos(phi))

U = (m11*(Xi-Xc)+m12*(Yi-Yc)+m13*(Zi-Zc))
V = (m21*(Xi-Xc)+m22*(Yi-Yc)+m23*(Zi-Zc))
W = (m31*(Xi-Xc)+m32*(Yi-Yc)+m33*(Zi-Zc))

xij = Xp - c * (U/W)
yij = Yp - c * (V/W)

def get_redundancy_numbers(A):
    # Compute hat matrix and redundancy numbers
    H = A @ np.linalg.inv(A.T @ A) @ A.T
    redundancy_numbers = 1 - np.diag(H)  # Redundancy numbers for each observation

    num_rows = len(A)
    num_cols = redundancy_numbers.size // num_rows
    redundancy_table = redundancy_numbers.reshape(num_rows, num_cols)

    s = np.sum(redundancy_numbers)
    print(f"Redundancy Numbers: {redundancy_numbers} Sum: {s}")
    return redundancy_numbers, s

# X, Y, Z, W, P, K
def get_deg_parameters(p):
    deg_parameters = copy.copy(p)
    deg_parameters[3] *= 180 / math.pi
    deg_parameters[4] *= 180 / math.pi
    deg_parameters[5] *= 180 / math.pi
    return deg_parameters

def output_solution(name, parameters, convert_to_degs, iterations, A, misclosure_vector, sigma_obs):
    if convert_to_degs:
        parameters = get_deg_parameters(parameters)

    print(f"Iterations: {iterations}")
    print(f"Estimated Parameters: {np.array(parameters)}")
    array_to_word_table(parameters, f"{name} Estimated Parameters")

    residuals = np.array(misclosure_vector)
    print(f"Residuals: {residuals}")
    array_to_word_table(residuals, f"{name} residuals")

    x_residuals = residuals[0::2]
    y_residuals = residuals[1::2]  
    array_to_word_table(x_residuals, f"{name} x_residuals")
    array_to_word_table(y_residuals, f"{name} y_residuals")
    rmsx = np.sqrt(np.mean(np.array(x_residuals)**2))
    rmsy = np.sqrt(np.mean(np.array(y_residuals)**2))

    print(f"RMS x: {rmsx}")
    print(f"RMS y: {rmsy}")

    r_numbers, r_sum = get_redundancy_numbers(A)
    x_redundancy = r_numbers[0::2]
    y_redundancy = r_numbers[1::2]
    x_redundancy = np.append(x_redundancy, np.sum(x_redundancy))
    y_redundancy = np.append(y_redundancy, np.sum(y_redundancy))
    array_to_word_table(x_redundancy, f"{name} x redundancy")
    array_to_word_table(y_redundancy, f"{name} y redundancy")

    P = (1 / sigma_obs**2) * np.eye(len(residuals))  # Weight matrix
    v = np.array(residuals)
    redundancy = len(v) - A.shape[1]  # degrees of freedom
    variance_factor = (v.T @ P @ v) / redundancy

    other_quantites = []
    other_quantites.append(sigma_obs)
    other_quantites.append(rmsx)
    other_quantites.append(rmsy)
    other_quantites.append(r_sum)
    other_quantites.append(variance_factor)
    print(f"other: {other_quantites}")
    array_to_word_table(other_quantites, f"{name} other quantities")

    # correlation matrix and standard deviation
    cov = sigma_obs**2 * np.linalg.inv(A.T @ A)
    std_devs = np.sqrt(np.diag(cov))  # Standard deviations of parameters
    corr_matrix = cov / np.outer(std_devs, std_devs)
    print("Correlation Matrix of Estimated Parameters:\n", corr_matrix)

    if convert_to_degs:
        std_devs = get_deg_parameters(std_devs)

    print(f"STD: {std_devs}")
    array_to_word_table(std_devs, f"{name} STD")
    array2d_to_word_table(corr_matrix.tolist(), f"{name} Correlation Matrix")

class Resection():
    def __init__(self, name, image, object, focal_length, std):
        self.name = name
        self.image = image
        self.object = object
        self.focal_length = focal_length
        self.std = std # mm

        self.partial_derivatives = []
        self.partial_derivatives.append(self.get_partial(xij))
        self.partial_derivatives.append(self.get_partial(yij))

        array2d_to_word_table(self.object, f"{self.name} object points")
        array2d_to_word_table(self.image, f"{self.name} image points")

        self.approximate_parameters()
        self.deg_parameters = self.get_deg_parameters(self.parameters)
        print(f"Approximated Parameters: {self.deg_parameters}")
        array_to_word_table(self.deg_parameters, f"{self.name} Approximated Parameters")

        self.estimate_parameters()

    def get_partial(self, o):
        row = []
        row.append(diff(o, Xc))
        row.append(diff(o, Yc))
        row.append(diff(o, Zc))
        row.append(diff(o, omega))
        row.append(diff(o, phi))
        row.append(diff(o, kappa))
        return row

    def get_deg_parameters(self, p):
        deg_parameters = copy.copy(p)
        deg_parameters[3] *= 180 / math.pi
        deg_parameters[4] *= 180 / math.pi
        deg_parameters[5] *= 180 / math.pi
        return deg_parameters

    def approximate_parameters(self):
        # Calculate centroids
        t_x = np.mean(self.object[:, 0])  # X centroid
        t_y = np.mean(self.object[:, 1])  # Y centroid
        x_mean = np.mean(self.image[:, 0])  # x centroid
        y_mean = np.mean(self.image[:, 1])  # y centroid

        # Center the coordinates
        X_centered = self.object[:, 0] - t_x
        Y_centered = self.object[:, 1] - t_y
        x_centered = self.image[:, 0] - x_mean
        y_centered = self.image[:, 1] - y_mean

        # Calculate the transformation parameters directly
        # These formulas approximate the least squares solution
        numerator_a = np.sum(X_centered * x_centered + Y_centered * y_centered)
        numerator_b = np.sum(Y_centered * x_centered - X_centered * y_centered)
        denominator = np.sum(x_centered**2 + y_centered**2)

        a = numerator_a / denominator
        b = numerator_b / denominator

        # Adjust the translation parameters
        t_x = t_x - a * x_mean + b * y_mean
        t_y = t_y - b * x_mean - a * y_mean

        # Calculate rotation angle
        theta_2D = np.arctan2(b, a)

        # Calculate scale factor
        lambda_approx = np.sqrt(a**2 + b**2)

        # Calculate Z coordinate of perspective center
        Z_ave = np.mean(self.object[:, 2])
        Z_C = self.focal_length * lambda_approx + Z_ave

        # Store the parameters
        self.parameters = [t_x, t_y, Z_C, 0, 0, theta_2D]

    def estimate_parameters(self):
        tol_coords = 0.012
        tol_tilt = 0.00056
        tol_k = 0.00053

        P = (1/self.std**2) * np.eye(len(self.image) * 2)

        iterations = 0
        while(1):
            A, misclosure_vector = self.design_matrix()

            corrections = -(np.linalg.inv(A.T @ P @ A) @ A.T @ P @ np.array(misclosure_vector))
            corrections = np.array(corrections).ravel()
            self.parameters = self.parameters + np.array(corrections).ravel()

            coords_corr = corrections[0:3]
            tilt_corr = corrections[3:5]
            k_corr = corrections[5]

            rel_corrections = np.abs(corrections / (np.abs(self.parameters) + 1e-10))
            iterations += 1
            if (np.max(np.abs(coords_corr)) < tol_coords and np.max(np.abs(tilt_corr)) < tol_tilt and np.abs(k_corr) < tol_k):
                A, misclosure_vector = self.design_matrix()
                output_solution(self.name, self.parameters, true, iterations, A, misclosure_vector, self.std)
                break

    def design_matrix(self):
        A = []
        misclosure_vector = []
        for point_index in range(len(self.image)):
            for coord_index in range(2):
                A_row = []
                
                for p in self.partial_derivatives[coord_index]:
                    value = self.evaluate(p, self.object[point_index])
                    A_row.append(value)

                A.append(A_row)
            misclosure_vector += self.misclosure(self.object[point_index], self.image[point_index])
        return np.matrix(A), misclosure_vector

    def misclosure(self, object_coord, image_coord):
        w = [
            self.evaluate(xij, object_coord) - image_coord[0],
            self.evaluate(yij, object_coord) - image_coord[1],
        ]
        return w

    def evaluate(self, eq, point):
        new_eq = eq.subs({
            Xc: self.parameters[0],
            Yc: self.parameters[1], 
            Zc: self.parameters[2], 
            omega: self.parameters[3], 
            phi: self.parameters[4], 
            kappa: self.parameters[5], 
            c: self.focal_length,
            Xi: point[0],
            Yi: point[1],
            Zi: point[2],
            Xp: 0,
            Yp: 0,
        })
        
        return float(N(new_eq, 10))
        
class Intersection():
    def evaluate(self, eq, eops, point):
        new_eq = eq.subs({
            Xc: eops[0],
            Yc: eops[1], 
            Zc: eops[2], 
            omega: eops[3], 
            phi: eops[4], 
            kappa: eops[5], 
            c: self.focal_length,
            Xi: point[0],
            Yi: point[1],
            Zi: point[2],
            Xp: 0,
            Yp: 0,
        })
        return float(N(new_eq, 50))

    def get_rad_parameters(self, p):
        dx = p[3] * math.pi / 180
        dy = p[4] * math.pi / 180
        dz = p[5] * math.pi / 180
        return [p[0], p[1], p[2], dx, dy, dz]

    def __init__(self, name, left_eops, right_eops, focal_length, std):
        self.name = name
        self.focal_length = focal_length
        self.left_eops = left_eops
        self.right_eops = right_eops
        self.std = std # mm

        left_c = np.array(left_eops[0:3])
        right_c = np.array(right_eops[0:3])
        baseline = np.linalg.norm(left_c - right_c)
        array_to_word_table([baseline], f"{self.name} baseline")

        array_to_word_table(self.left_eops, f"{self.name} left eops")
        array_to_word_table(self.right_eops, f"{self.name} right eops")

        self.left_eops = self.get_rad_parameters(self.left_eops)
        self.right_eops = self.get_rad_parameters(self.right_eops)

        self.partial_derivatives = []
        self.partial_derivatives.append(self.get_partial(xij))
        self.partial_derivatives.append(self.get_partial(yij))

    def evaluate_b(self, eq, eops, image):
        new_eq = eq.subs({
            Xc:  eops[0],
            Yc:  eops[1], 
            Zc:  eops[2], 
            omega:  eops[3], 
            phi:  eops[4], 
            kappa:  eops[5], 
            c: self.focal_length,
            xij_bar: image[0],
            yij_bar: image[1]
        })
        return np.float64(N(new_eq, 50))

    def approximate_parameters(self):
        a11 = xij_bar*m31 + c*m11
        a12 = xij_bar*m32 + c*m12
        a13 = xij_bar*m33 + c*m13

        a21 = yij_bar*m31 + c*m21
        a22 = yij_bar*m32 + c*m22
        a23 = yij_bar*m33 + c*m23

        eq_A = np.matrix([
            [a11, a12, a13],
            [a21, a22, a23],
            [a11, a12, a13],
            [a21, a22, a23],
        ])

        A = np.zeros((4,3))

        for i in range(3): A[0, i] = self.evaluate_b(eq_A[0, i], self.left_eops, [self.image[0], self.image[1]])
        for i in range(3): A[1, i] = self.evaluate_b(eq_A[1, i], self.left_eops, [self.image[0], self.image[1]])
        for i in range(3): A[2, i] = self.evaluate_b(eq_A[2, i], self.right_eops, [self.image[2], self.image[3]])
        for i in range(3): A[3, i] = self.evaluate_b(eq_A[3, i], self.right_eops, [self.image[2], self.image[3]])

        b1 = (xij_bar*m31+c*m11)*Xc + (xij_bar*m32+c*m12)*Yc + (xij_bar*m33+c*m13)*Zc
        b2 = (yij_bar*m31+c*m21)*Xc + (yij_bar*m32+c*m22)*Yc + (yij_bar*m33+c*m23)*Zc

        b = [
            self.evaluate_b(b1, self.left_eops, [self.image[0], self.image[1]]),
            self.evaluate_b(b2, self.left_eops, [self.image[0], self.image[1]]),
            self.evaluate_b(b1, self.right_eops, [self.image[2], self.image[3]]),
            self.evaluate_b(b2, self.right_eops, [self.image[2], self.image[3]])
        ]

        return np.linalg.inv(A.T @ A) @ A.T @ b

    def do_iterations(self, name, left, right, output):
        self.name = name
        self.image = [left[0], left[1], right[0], right[1]]
        parameters = self.approximate_parameters()
        A, w = self.design_matrix(parameters)

        if output:
            print(self.image)
            array_to_word_table(self.image, f"{self.name} image points")
            print(parameters)
            array_to_word_table(parameters, f"{self.name} approximate parameters")      
            print(A)
            print(w)

        iterations = 0
        while(1):
            A, misclosure_vector = self.design_matrix(parameters)

            corrections = -(np.linalg.inv(A.T @ A) @ A.T @ np.array(misclosure_vector))
            parameters = parameters + np.array(corrections).ravel()

            rel_corrections = np.abs(corrections / (np.abs(parameters) + 1e-10))
            iterations += 1
            if np.max(rel_corrections) < 1e-3: # if converged
                A, misclosure_vector = self.design_matrix(parameters)
                if output:
                    output_solution(self.name, parameters, false, iterations, A, misclosure_vector, self.std)
                break

        return parameters

    def get_partial(self, o):
        row = []
        row.append(diff(o, Xc))
        row.append(diff(o, Yc))
        row.append(diff(o, Zc))
        return row
    
    def design_matrix(self, parameters):
        A = []
        misclosure_vector = []
        for coord_index in range(4):
            if coord_index == 0 or coord_index == 1:
                eops = self.left_eops
            else:
                eops = self.right_eops

            if coord_index % 2 == 0: # x
                eq = xij
                p_e_index = 0
            else: # y
                eq = yij
                p_e_index = 1
            p_e = self.partial_derivatives[p_e_index]

            A_row = []
            for p in p_e:
                value = -self.evaluate(p, eops, parameters)
                A_row.append(value)

            A.append(A_row)
            misclosure_vector += [self.evaluate(eq, eops, parameters) - self.image[coord_index]]

        return np.matrix(A), misclosure_vector

def do_resection_test():
    print("\n\nResection Validation\n")
    # test data 
    test_resection_image_points = np.array([
        [106.399, 90.426], # 30
        [18.989,  93.365],  # 40
        [98.681, -62.769], # 50
        [9.278,  -92.926]   # 112
    ]) # mm (reduced to PP)

    test_resection_object_points = np.array([
        [7350.27, 4382.54, 276.42],
        [6717.22, 4626.41, 280.05],
        [6905.26, 3279.84, 266.47],
        [6172.84, 3269.45, 248.10]
    ]) # m

    test_res = Resection("test res", test_resection_image_points, test_resection_object_points, 152.150, 0.015)

def do_single_photo_resection():
    print("\n\nResection\n")
    # from lab 2
    image_27 = np.array([
        [-9.590463762012142, 96.21844317125414],
        [-2.4130665595082577, -5.994692696706493],
        [18.94331818098372, -81.81523920233901],
        [90.37866498271781, -91.09240982867657],
        [18.149336968374765, 109.57529999437965],
        [44.59762997126615, 7.472924495358117],
        [-7.65703502813932, -49.11171271008798],
        [52.69132378762902, -93.1781834642561],
    ])

    image_28 = np.array([
        [-105.46911858501092, 98.73604528901996],
        [-95.08144734697845, -4.847626634342134],
        [-72.54716548403219, -79.76433289820899],
        [-1.3569472252226806, -86.94970555134802],
        [-77.82635997556542, 113.4048797093696],
        [-48.846121573185016, 10.130960551399694],
        [-98.85497251911589, -48.0676549863933],
        [-38.93560060084638, -90.07943042220953],
    ])

    control_indices = [2, 3, 4, 6] # what indices are control points
    gcps = object_coords[control_indices]
    image_coords_27 = image_27[control_indices]
    image_coords_28 = image_28[control_indices]

    res_27 = Resection("image 27 res", image_coords_27, gcps, 153.358, images[0].rms)
    res_28 = Resection("image 28 res", image_coords_28, gcps, 153.358, images[1].rms)

    do_resection_test()

    return res_27, res_28

def do_intersection_test():
    print("\n\nIntersection Validation\n")
    test_inter_image_points_left = np.array([
        [70.964, 4.907],  # 72
        [-0.931, -7.284], # 127
    ])

    test_inter_image_points_right = np.array([
        [-15.581, -0.387],  # 72
         [-85.407, -8.351], # 127
    ])

    # [Left, Right]
    test_intersection_eops = np.array([
        [6349.488,	7021.897], # Xc (m)		
        [3965.252,	3775.680], # Yc (m)	
        [1458.095,	1466.702], # Zc (m)	
        [0.9885,	1.8734],   # w (dd)		
        [0.4071,	1.6751],   # p (dd)	
        [-18.9049,	-15.7481], # k (dd)	
    ])

    test_inter = Intersection("test intersection", test_intersection_eops[:, 0], test_intersection_eops[:, 1], 152.150, 0.015)
    test_inter.do_iterations("test intersection 72", test_inter_image_points_left[0], test_inter_image_points_right[0], true)
    test_inter.do_iterations("test intersection 127", test_inter_image_points_left[1], test_inter_image_points_right[1], true)

def do_space_intersection(res_27, res_28):
    print("\n\nIntersection\n")
    # measurements
    quinn_27 = [
        [15334, 3241], [15539, 3240], [16007, 3241], [16250, 3241],
        [15337, 3392], [15577, 3393], [16004, 3392], [16244, 3392],
        [15338, 3582], [15577, 3582], [16004, 3583], [16244, 3582],
        [15338, 3792], [15577, 3793], [16005, 3793], [16245, 3794]
    ] # px

    fre_27 = [
        [15333, 3241], [15539, 3241], [16007, 3240], [16251, 3243],
        [15338, 3392], [15578, 3392], [16003, 3392], [16242, 3392],
        [15336, 3582], [15577, 3582], [16005, 3582], [16245, 3582],
        [15337, 3791], [15579, 3790], [16005, 3791], [16244, 3792]
    ] # px

    quinn_28 = [
        [7365, 2859], [7574, 2855], [8049, 2839], [8294, 2828],
        [7373, 3017], [7615, 3008], [8048, 2996], [8290, 2988],
        [7378, 3211], [7620, 3204], [8054, 3189], [8295, 3181],
        [7383, 3427], [7625, 3419], [8058, 3406], [8301, 3399]
    ] # px

    fre_28 = [
        [7366, 2861], [7575, 2856], [8049, 2839], [8295, 2826],
        [7374, 3017], [7616, 3010], [8047, 2996], [8288, 2988],
        [7377, 3211], [7620, 3203], [8053, 3190], [8296, 3182],
        [7383, 3426], [7625, 3419], [8057, 3406], [8301, 3398]
    ] # px

    combined_27 = np.array([[[q[0], q[1]], [f[0], f[1]]] for q, f in zip(quinn_27, fre_27)])
    combined_28 = np.array([[[q[0], q[1]], [f[0], f[1]]] for q, f in zip(quinn_28, fre_28)])

    combined_27 = images[0].pool_data(combined_27)
    combined_28 = images[1].pool_data(combined_28)

    images[0].correct(combined_27)
    images[1].correct(combined_28)

    array2d_to_word_table(combined_27, f"space intersection image 27 measurements")
    array2d_to_word_table(combined_28, f"space intersection image 28 measurements")

    lab_inter = Intersection("lab intersection", res_27.deg_parameters, res_28.deg_parameters, 153.358, np.sqrt(images[0].rms**2 + images[1].rms**2))
    object_coords = np.zeros((len(combined_27), 3))
    for i in range(len(combined_27)):
        object_coords[i] = lab_inter.do_iterations("lab intersection", combined_27[i], combined_28[i], false)

    array2d_to_word_table(object_coords, f"space intersection object coords")

    mean_height = np.mean(object_coords[:, 2])
    residuals = object_coords[:, 2] - mean_height
    RMSE = np.sqrt(np.mean(residuals**2))

    stats = [ mean_height, RMSE ]
    array_to_word_table(stats, f"space intersection stats")
    array_to_word_table(residuals, f"space intersection object coords z residuals")

    print(f"mean_height: + {mean_height}")
    print(f"residuals: {residuals}")
    print(f"RMSE: {RMSE}")

    x_coords = object_coords[:, 0]
    y_coords = object_coords[:, 1]

    # We can treat the residuals as vectors in the z-direction (height residuals)
    plt.figure(figsize=(10, 8))
    plt.quiver(x_coords, y_coords, np.zeros_like(residuals), residuals, angles='xy', scale_units='xy', scale=0.1, color='red', width=0.005)

    # Add labels and title
    plt.title('Quiver Plot of Height Residuals')
    plt.xlabel('X (m)')
    plt.ylabel('Y (m)')

    # Show the plot
    plt.grid(True)
    plt.show()

    do_intersection_test()

def main():
    res_27, res_28 = do_single_photo_resection()
    do_space_intersection(res_27, res_28)

if __name__ == '__main__':
    main()
    doc.save("output.docx")