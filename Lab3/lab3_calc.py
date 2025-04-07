from sympy import *
import numpy as np
import copy
import math
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document

doc = Document()

# print a 2d array to a word table
def array_to_word_table(array, name, value_type, decimals=2):
    doc.add_paragraph(name)
    table = doc.add_table(rows=len(array), cols=len(array[0]))

    for i, row in enumerate(array):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(round(cell, decimals))

# Notes:
# 1 pixel for Y-Parallax

# create equations
omega, phi, kappa = symbols('omega phi kappa')

r1 = Matrix([
    [1, 0, 0],
    [0, cos(omega), -sin(omega)],
    [0, sin(omega), cos(omega)],
])

r2 = Matrix([
    [cos(phi), 0, sin(phi)],
    [0, 1, 0],
    [-sin(phi), 0, cos(phi)],
])

r3 = Matrix([
    [cos(kappa), -sin(kappa), 0],
    [sin(kappa), cos(kappa), 0],
    [0, 0, 1],
])

# Create rotation matrix elements
rotation_matrix = r1 @ r2 @ r3

class Relative_Orientation:
    bx, by, bz, l1, l2, r1, r2, f = symbols('bx by bz l1 l2 r1 r2 f')

    right_image = Matrix([r1, r2, -f])
    right_prime = rotation_matrix @ right_image

    # Misclosure equation
    mis = Matrix([
        [bx, by, bz],
        [l1, l2, -f],
        [right_prime[0], right_prime[1], right_prime[2]]
    ])

    mis_det = mis.det()

    # Partial derivates for design matrix
    partial_derivatives = []
    partial_derivatives.append(diff(mis_det, by))
    partial_derivatives.append(diff(mis_det, bz))
    partial_derivatives.append(diff(mis_det, omega))
    partial_derivatives.append(diff(mis_det, phi))
    partial_derivatives.append(diff(mis_det, kappa))

    # Unique scale factors
    rx, ry, rz, lx, ly, lz = symbols('rx ry rz lx ly lz')
    lambda_i = (bx*rz-bz*rx) / (lx*rz+f*rx)
    mu_i = (-bx*f-bz*lx) / (lx*rz+f*rx)

    # from last lab
    pixel_spacing = 0.011901151249225804

    def __init__(self, tie_left, tie_right, baseline, focal_length):
        self.tie_left = tie_left
        self.tie_right = tie_right
        self.baseline = baseline
        self.focal_length = focal_length

        #self.plot()

        self.parameters = np.array([0.0, 0.0, 0.0, 0.0, 0.0]) # by, bz, omega, phi, kappa
        self.estimate_parameters()
        self.set_deg_parameters()
        print(f"Final Relative Parameters: {self.deg_parameters}")
        self.right_pc = [self.baseline, self.parameters[0], self.parameters[1]]
        self.rotation_parameters = [self.parameters[2], self.parameters[3], self.parameters[4]]

    def evaluate(self, eq, left_coord, right_coord):
        new_eq = eq.subs({
            self.bx: self.baseline,
            self.f: self.focal_length,
            self.l1: left_coord[0],
            self.l2: left_coord[1],
            self.r1: right_coord[0],
            self.r2: right_coord[1],
            self.by: self.parameters[0],
            self.bz: self.parameters[1],
            omega: self.parameters[2],
            phi: self.parameters[3],
            kappa: self.parameters[4]
        })
        
        return float(N(new_eq, 50))

    def design_matrix(self):
        A = []
        misclosure_vector = []
        
        for i in range(len(self.tie_left)):
            A_row = []
            for p in self.partial_derivatives:
                value = self.evaluate(p, self.tie_left[i], self.tie_right[i])
                A_row.append(float(value))
            A.append(A_row)

            misclosure_vector.append(self.evaluate(self.mis_det, self.tie_left[i], self.tie_right[i]))

        return np.matrix(A), misclosure_vector

    def estimate_parameters(self):
        iterations = 0
        while(1):
            A, misclosure_vector = self.design_matrix()

            corrections = -(np.linalg.inv(A.T @ A) @ A.T @ np.array(misclosure_vector))
            self.parameters = self.parameters + np.array(corrections).ravel()
            iterations = iterations + 1

            rel_corrections = np.abs(corrections / (np.abs(self.parameters) + 1e-10))
            if np.max(rel_corrections) < 1e-10:
                A, misclosure_vector = self.design_matrix()

                cov = np.linalg.inv(A.T @ A)
                std_devs = np.sqrt(np.diag(cov))  # Standard deviations of parameters
                corr_matrix = cov / np.outer(std_devs, std_devs)
                print("Correlation Matrix of Estimated Parameters:\n", corr_matrix)
                #array_to_word_table(corr_matrix.tolist(), "Correlation Matrix", float, decimals=4)
                break

    def set_deg_parameters(self):
        self.deg_parameters = copy.copy(self.parameters)
        self.deg_parameters[2] *= 180.0 / np.pi
        self.deg_parameters[3] *= 180.0 / math.pi
        self.deg_parameters[4] *= 180.0 / math.pi

    def plot(self):
        left = np.array(self.tie_left)
        right = np.array(self.tie_right)

        plt.figure(figsize=(10, 5))

        # Subplot for Image 27
        plt.subplot(1, 2, 1)
        plt.plot(left[:, 0], left[:, 1], 'ro')
        plt.title('Image 27 Tie Points')
        plt.xlabel('x_L (mm)')
        plt.ylabel('y_L (mm)')
        plt.grid(True)
        plt.axis([-125, 125, -125, 125])

        # Subplot for Image 28
        plt.subplot(1, 2, 2)
        plt.plot(right[:, 0], right[:, 1], 'bo')
        plt.title('Image 28 Tie Points')
        plt.xlabel('x_R (mm)')
        plt.ylabel('y_R (mm)')
        plt.grid(True)
        plt.axis([-125, 125, -125, 125])

        plt.tight_layout()
        plt.show()

    def space_intersection_point(self, left, right):
        l_i = N(self.lambda_i.subs({
            self.bx: self.right_pc[0],
            self.bz: self.right_pc[2],
            self.f: self.focal_length,
            self.lx: left[0],
            self.rx: right[0],
            self.rz: right[2]
        }))

        m_i = N(self.mu_i.subs({
            self.bx: self.right_pc[0],
            self.bz: self.right_pc[2],
            self.f: self.focal_length,
            self.lx: left[0],
            self.rx: right[0],
            self.rz: right[2]
        }))

        new_lx =  l_i * left[0]
        new_ly =  l_i * left[1]
        new_lz = -l_i * self.focal_length

        new_rx = m_i * right[0] + self.right_pc[0]
        new_ry = m_i * right[1] + self.right_pc[1]
        new_rz = m_i * right[2] + self.right_pc[2]

        # Calculating Y_Parallax
        y_parallax = new_ry - new_ly

        if not math.isclose(new_lx, new_rx): print(f"XmL ({new_lx}) != XmR ({new_rx})")
        mean_y = (new_ly + new_ry) / 2
        if not math.isclose(new_lz, new_rz): print(f"ZmL ({new_lz}) != ZmR ({new_rz})")

        return [new_lx, mean_y, new_lz], y_parallax

    def space_intersection(self, left, right):
        # apply the transformation matrix to the right image points
        model_right = []
        for i in range(len(right)):
            result = N(self.right_prime.subs({
                self.f: self.focal_length,
                self.r1: right[i][0],
                self.r2: right[i][1],
                self.bx: self.baseline,
                self.by: self.parameters[0],
                self.bz: self.parameters[1],
                omega: self.parameters[2],
                phi: self.parameters[3],
                kappa: self.parameters[4]
            }))
            model_right.append([float(result[0]), float(result[1]), float(result[2])])

        y_parallaxs = []
        model_coordinates = []
        for i in range(len(left)):
            coordinates, y_parallax = self.space_intersection_point(left[i], model_right[i])
            model_coordinates.append(coordinates)
            y_parallaxs.append(y_parallax)

        array_to_word_table([y_parallaxs], "Y Parallax", float, decimals=4)

        print(f"Y Parallaxs {y_parallaxs}")
        print(f"Pixel Y Parallaxs {np.array(y_parallaxs) * (1/self.pixel_spacing)}")
        print(f"Model Coordinates {model_coordinates}")
        return model_coordinates

def main():
    class Data:
        def __init__(self):
            self.control = []
            self.tie = []
            self.index_map = {}

    # Define the data

    # original data
    '''
    data = {
        "ID": [1, 2, 3, 4, 5, 6, 7, 8, 100, 102, 104, 105, 200, 201, 202, 203, "T1", "T2", "T3", "T4", "T5", "T6"],
        "x1": [-106, 106.004, -106.002, 106.017, -111.995, 112.003, 0.011, 0.004, -9.59, -2.413, 18.943, 90.379, 18.149, 44.598, -7.657, 52.691, -10.105, 94.369, -10.762, 90.075, -9.489, 85.42],
        "y1": [-106.004, 105.999, 106.006, -105.982, -0.002, -0.018, 112.01, -112.01, 96.218, -5.995, -81.815, -91.092, 109.575, 7.473, -49.112, -93.178, 15.011, -4.092, -104.711, -91.378, 96.26, 103.371],
        "x2": [-105.991, 106.011, -106.011, 106.004, -111.999, 112.003, 0.015, 0.01, -105.469, -95.081, -72.547, -1.357, -77.826, -48.846, -98.855, -38.936, -103.829, 0.868, -100.169, -1.607, -105.395, -9.738],
        "y2": [-105.995, 106.006, 105.996, -106.003, -0.003, -0.006, 112.007, -112.005, 98.736, -4.848, -79.764, -86.95, 113.405, 10.131, -48.068, -90.079, 16.042, -0.022, -103.14, -87.253, 98.706, 109.306]
    }
    '''

    # more precise data
    #'''
    data = {
        "ID": [1, 2, 3, 4, 5, 6, 7, 8, 100, 102, 104, 105, 200, 201, 202, 203, "T1", "T2", "T3", "T4", "T5", "T6"],
        "x1": [
            -105.99978691496106, 106.0037749724548, -106.00171609975594, 106.01650056099524,
            -111.99473580642875, 112.00312152191853, 0.010796402230584446, 0.004339112954062496,
            -9.590463762012142, -2.4130665595082577, 18.94331818098372, 90.37866498271781,
            18.149336968374765, 44.59762997126615, -7.65703502813932, 52.69132378762902,
            -10.104572641422829, 94.36910810514888, -10.762476440703479, 90.0752611615937,
            -9.489320272082233, 85.41964342146787
        ],
        "y1": [
            -106.00441822465683, 105.9987082689118, 106.00568167216062, -105.98183242323235,
            -0.0020832945284820097, -0.01781006222153149, 112.01048519203304, -112.01016240128253,
            96.21844317125414, -5.994692696706493, -81.81523920233901, -91.09240982867657,
            109.57529999437965, 7.472924495358117, -49.11171271008798, -93.1781834642561,
            15.010873181699598, -4.091842390658864, -104.71141379376125, -91.37807702186743,
            96.2601088526174, 103.3708182671729
        ],
        "x2": [
            -105.99076409478431, 106.01110325284907, -106.01105312699268, 106.00352653557252,
            -111.99948724156867, 112.00341404148837, 0.015083026729605476, 0.010463493049818872,
            -105.46911858501092, -95.08144734697845, -72.54716548403219, -1.3569472252226806,
            -77.82635997556542, -48.846121573185016, -98.85497251911589, -38.93560060084638,
            -103.82924974464979, 0.868221431682023, -100.16923346397878, -1.6066403254265023,
            -105.39473279961206, -9.738129753479935
        ],
        "y2": [
            -105.99486652288773, 106.00626660936378, 105.99603184454695, -106.00269050440353,
            -0.0027548768451323893, -0.005981950552230421, 112.00724730255126, -112.00468738098284,
            98.73604528901996, -4.847626634342134, -79.76433289820899, -86.94970555134802,
            113.4048797093696, 10.130960551399694, -48.0676549863933, -90.07943042220953,
            16.041503137295727, -0.02152744389137064, -103.14023006192618, -87.25335975412447,
            98.70635813160811, 109.30613201141969
        ]
    }
    #'''
    df = pd.DataFrame(data)

    # Define control point IDs
    control_points = {100, 102, 104, 105, 200, 201, 202, 203}

    image_27 = Data()
    image_28 = Data()

    for index, row in df.iterrows():
        point_id, x27, y27, x28, y28 = row.to_list()

        if isinstance(point_id, int) and point_id in control_points:
            # It's a control point
            image_27.index_map[point_id] = len(image_27.control)
            image_28.index_map[point_id] = len(image_28.control)

            image_27.control.append((x27, y27))
            image_28.control.append((x28, y28))
        elif isinstance(point_id, str) and point_id.startswith("T"):
            # It's a tie point
            image_27.index_map[point_id] = len(image_27.tie)
            image_28.index_map[point_id] = len(image_28.tie)

            image_27.tie.append((x27, y27))
            image_28.tie.append((x28, y28))

    #plot_points(image_27.tie, image_28.tie)

    print("\n\nLab Data\n")
    print(image_27.tie)
    lab_rel = Relative_Orientation(image_27.tie, image_28.tie, 92, 153.358)

    array_to_word_table([lab_rel.deg_parameters], "Parameters", float, decimals=4)

    tie_points = lab_rel.space_intersection(image_27.tie, image_28.tie)
    array_to_word_table(tie_points, "Tie Points", float, decimals=4)

    control_points = lab_rel.space_intersection(image_27.control, image_28.control)
    array_to_word_table(control_points, "Control Points", float, decimals=4)

    # Verification
    print("\n\nTest Data\n")

    test_data_left = [
        [106.399, 90.426], # 30
        [18.989, 93.365],  # 40
        [70.964, 4.907],   # 72
        [-0.931, -7.284],  # 127
        [9.278, -92.926],  # 112
        [98.681, -62.769]  # 50
    ]

    test_data_right = [
        [24.848, 81.824],  # 30
        [-59.653, 88.138], # 40
        [-15.581, -0.387], # 72 
        [-85.407, -8.351], # 127
        [-78.81, -92.62],  # 112
        [8.492, -68.873]   # 50
    ]

    # test data parameters
    test_rel = Relative_Orientation(test_data_left, test_data_right, 92, 152.150)
    array_to_word_table([test_rel.deg_parameters], "Test Parameters", float, decimals=4)

    # used to test if the space intersection is working
    #target_parameters = [5.0455, 2.1725, math.radians(0.4392),  math.radians(1.508), math.radians(3.1575)]

    test_points = test_rel.space_intersection(test_data_left, test_data_right)
    array_to_word_table(test_points, "Test Points", float, decimals=4)

    # printing answers out to a word document
    doc.save("output.docx")

if __name__ == "__main__":
    main()