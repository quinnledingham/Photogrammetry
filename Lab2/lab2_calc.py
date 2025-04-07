import pandas as pd
import math
import numpy as np
import matplotlib.pyplot as plt
import copy

from docx import Document

object_coords = [
    [-399.28, -679.72, 1090.96],
    [109.70, -642.35, 1086.43],
    [475.55, -538.18, 1090.50],
    [517.62, -194.43, 1090.65],
    [-466.39, -542.31, 1091.55],
    [42.73, -412.19, 1090.82],
    [321.09, -667.45, 1083.49],
    [527.78, -375.72, 1092.00]
]

# 4 observations for each point
data = np.array([
    #   x1     y1     x2     y2     x3     y3
    # image 27
    [ 1347, 19285,  1347, 19285,  1347, 19284,  1347, 19285], # F1
    [19161,  1470, 19175,  1484, 19160,  1469, 19162,  1470], # F2
    [ 1345,  1471,  1358,  1473,  1345,  1471,  1346,  1472], # F3
    [19163, 19283, 19163, 19283, 19167, 19282, 19163, 19283], # F4
    [  843, 10379,   849, 10378,   842,  10378,  843, 10379], # F5
    [19666, 10376, 19672, 10391, 19666, 10377, 19666, 10377], # F6
    [10254,   966, 10267,   973, 10253,  966,  10254,   967], # F7
    [10254, 19789, 10255, 19789, 10255, 19789, 10255, 19789], # F8
    [ 1347, 19285,  1347, 19286,  1348, 19286,  1347, 19285], # 100
    [19175,  1483, 19175,  1484, 19175,  1481, 19175,  1483], # 102
    [ 1358,  1472,  1358,  1472,  1358,  1472,  1358,  1472], # 104
    [19162, 19297, 19162, 19297, 19161, 19296, 19163, 19297], # 105
    [  849,  10379,  849, 10378,   848,  10379,  849, 10379], # 200
    [19672, 10391, 19672, 10391, 19672, 10390, 19672, 10391], # 201
    [10267,   972, 10268,   973, 10268,   974, 10268,   973], # 202
    [10254, 19795, 10254, 19796, 10254, 19796, 10255, 19795], # 203
    [ 9447,  2295,  9458,  2293,  9448,  2294,  9447,  2296], # T1
    [10053, 10883, 10056, 10883, 10048, 10880, 10052, 10883], # T2
    [11843, 17255, 11851, 17259, 11848, 17245, 11844, 17251], # T3
    [17844, 18028, 17852, 18033, 17862, 18039, 17841, 18027], # T4
    [11782,  1172, 11781,  1172, 11781,  1172, 11781,  1174], # T5
    [14000,  9752, 14008,  9749, 14007,  9748, 13998,  9754], # T6
    # image 28
    [ 9612, 14506,  9611, 14507,  9612, 14502,  9610, 14505], # F1
    [14684, 18207, 14676, 18211, 14685, 18205, 14685, 18205], # F2
    [ 1404,  2083,  1408,  2080,  1402,  2081,  1398,  2084], # F3
    [ 2269, 10787,  2272, 10787,  2270, 10787,  2268, 10786], # F4
    [ 4158, 17086,  4160, 17080,  4159, 17083,  4159, 17083], # F5
    [10140, 17689, 10147, 17696, 10140, 17689, 10136, 17688], # F6
    [ 3727,   851,  3726,   849,  3727,   852,  3727,   852], # F7
    [ 6154,  9531,  6163,  9527,  6153,  9532,  6152,  9532], # F8
    [ 1948, 14420,  1958, 14415,  1947, 14419,  1948, 14419], # 100
    [ 6982, 17952,  6982, 17955,  6983, 17950,  6984, 17949], # 102
    [ 9406,  9117,  9405,  9118,  9407,  9117,  9406,  9117], # 104
    [18185, 10723, 18186, 10722, 18187, 10722, 18185, 10723], # 105
    [ 9350, 19175,  9350, 19176,  9350, 19177,  9350, 19176], # 200
    [17824, 18055, 17824, 18056, 17825, 18056, 17824, 18056], # 201
    [ 9459,  2291,  9458,  2291,  9458,  2291,  9459,  2291], # 202
    [17434,  1693, 17435,  1694, 17435,  1695, 17434,  1694], # 203
    [ 1537,  9029,  1536,  9034,  1535,  9032,  1536,  9029], # T1
    [10332, 10387, 10334, 10384, 10333, 10387, 10332, 10387], # T2
    [ 1836, 19046,  1838, 19046,  1836, 19046,  1836, 19046], # T3
    [10120, 17716, 10118, 17714, 10121, 17717, 10120, 17717], # T4
    [ 1410,  2080,  1408,  2095,  1409,  2082,  1410,  2081], # T5
    [ 9448,  1199,  9445,  1200,  9450,  1199,  9449,  1199]  # T6
])

'''
def dist_3D(a, b):
    return math.sqrt((a[0] - b[0])**2 + (a[1] - b[1])**2 + (a[2] - b[2])**2)

longest_dist = 0
for a in object_coords:
    for b in object_coords:
        if a is b:
            continue
        dist = dist_3D(a, b)
        if dist > longest_dist:
            longest_dist = dist
            print(f"A: {a}, B: {b}")
'''

def load_data(data, low, high):
    load = []
    for i in range(low, high):
        same = []
        for j, value in enumerate(data[i]):
            if j % 2 == 0:
                coords = [data[i][j], data[i][j+1]]
                same.append(coords)
        load.append(same)
    return load

def right_hand(a):
    for coords in a:
        x, y = coords
        #x = coords[0] - (image_dim[0]/2)
        #y = (image_dim[1]/2) - coords[1]

        coords[0] = x
        coords[1] = -y


def mean(a):
    sum = [0, 0]
    for v in a:
        if not math.isnan(v[0]):
            sum[0] = sum[0] + v[0]
        if not math.isnan(v[1]):
            sum[1] = sum[1] + v[1]
    sum[0] = sum[0] / len(a)
    sum[1] = sum[1] / len(a)
    return sum

def std(a):
    m = mean(a)
    sum = [0, 0]
    for v in a:
        if not math.isnan(v[0]):
            sum[0] = sum[0] + (v[0] - m[0])**2
        if not math.isnan(v[1]):
            sum[1] = sum[1] + (v[1] - m[1])**2
    sum[0] = math.sqrt(sum[0]/(len(a) - 1))
    sum[1] = math.sqrt(sum[1]/(len(a) - 1))
    return sum

def array_3d_to_2d(a):
    output = []

    for v1 in a:
        e = []
        for v2 in v1:
            for v3 in v2:
                e.append(v3)
        output.append(e)

    return output

def concat_2d(a, b):
    return [row1 + row2 for row1, row2 in zip(a, b)]

# print a 2d array to a word table
def array_to_word_table(array, name, type, decimals=2):
    doc.add_paragraph(name)
    table = doc.add_table(rows=len(array), cols=len(array[0]))

    for i, row in enumerate(array):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(type(round(cell, decimals)))

def design_matrix(fiducial):
    n = len(fiducial)
    A = np.zeros((2 * n, 6))

    for i in range(n):
        x, y = fiducial[i]
        A[2*i, 0] = x
        A[2*i, 1] = y
        A[2*i, 2] = 1
        A[2*i + 1, 3] = x
        A[2*i + 1, 4] = y
        A[2*i + 1, 5] = 1

    return A

def residual_plot(og, residuals):
    print("\n\nPlot")

    og = np.array(og)
    r = np.array(residuals)

    plt.figure(figsize=(8, 6))

    for i in range(len(og)):
        plt.arrow(og[i, 0], og[i, 1], 1000 * r[i, 0], 1000 * r[i, 1], color='red', head_width=1)

    plt.xlabel("x (mm)")
    plt.ylabel("y (mm)")
    #plt.axhline(y=0, color='black', linewidth=0.7, linestyle='--')
    plt.legend()
    plt.show()

def RMS(a):
    sum = 0
    for v in a:
        sum = sum + v**2
    
    return math.sqrt(sum/len(a))

def affine_transformation(data, x_hat):
    for coords in data:
        x = (x_hat[0] * coords[0] + x_hat[1] * coords[1]) + x_hat[2]
        y = (x_hat[3] * coords[0] + x_hat[4] * coords[1]) + x_hat[5]

        coords[0] = x
        coords[1] = y

# data is the fiducials in a [[]]

def get_affine_transformation_parameters(data, t_data, tag):
    print(f"\n{tag}:")
    print(data)

    l = []
    for coords in t_data:
        l.append(coords[0])
        l.append(coords[1])

    l = np.array(l)
    A = design_matrix(data)

    x_hat = np.linalg.inv(A.T @ A) @ A.T @ l
    print(f"x_hat: {x_hat}")
    # a b x c d y

    # displaying residuals in stats and graphs

    residuals = A @ x_hat - l
    print(f"residuals: {residuals}")

    # convert 1D-array into 2D-array
    r = []
    r_3d = []
    for i in range(len(residuals)):
        if i % 2 == 0:
            coord = [residuals[i], residuals[i + 1]]
            r.append(coord)
            r_3d.append([coord])

    #residual_plot(t_data, r)

    r = np.array(r)
    print(f"RMS x: {RMS(r[:, 0])}, y: {RMS(r[:, 1])}")

    array_to_word_table(r, f"residual_{tag}", float, decimals=6)

    return x_hat

class Camera:
    def __init__(self):
        self.ppo = [-0.006, 0.006] # mm

        self.k0 = 0.8878e-4
        self.k1 = -0.1528e-7
        self.k2 = 0.5256e-12
        self.k3 = 0

        self.p1 = 0.1346e-6
        self.p2 = 0.1224e-7

        self.f = 153.358 # focal length (mm)

        self.fiducial_marks = [
            [-105.997, -105.995],
            [106.004, 106.008],
            [-106.000, 106.009],
            [106.012, -105.995],
            [-112.000, 0.007],
            [112.006, 0.007],
            [0.005, 112.007],
            [0.002, -111.998]
        ]

class Image:
    def __init__(self, camera, tag, dim_x, dim_y, data):
        self.camera = camera
        self.tag = tag
        self.dim = [dim_x, dim_y]
        self.all_data = data

    def pool_data(self, in_data):
        for v in in_data:
            right_hand(v)

        values = []
        table = []
        for v in in_data:
            m = mean(v)
            s = std(v)
            values.append(m)
            table.append([m[0], m[1], s[0], s[1]])

        return values
        
    # combines all of the individual measurements
    def pool(self):
        self.data = {key: 0 for key in self.all_data.keys()}

        # applying right-handed transformation
        for key in self.all_data:
            for v in self.all_data[key]:
                right_hand(v)

            array_to_word_table(array_3d_to_2d(self.all_data[key]), f"{key}_{self.tag}", int)
            print(f"{key}_{self.tag}\n{self.all_data[key]}")

        # calculating mean and std
        for i, key in enumerate(self.all_data):
            a = self.all_data[key]
            values = []

            table = []

            print(f"{key}:")
            for v in a:
                m = mean(v)
                s = std(v)
                print(f"mean: {m}, std: {s}")
                values.append(m)
                table.append([m[0], m[1], s[0], s[1]])

            self.data[key] = values

            array_to_word_table(table, f"mean__std_{key}_{self.tag}", float)

    def get_affine_transformation_parameters(self):
        self.x_hat = get_affine_transformation_parameters(self.data['fiducials'], self.camera.fiducial_marks, self.tag)

    def apply_transformation(self):
        for key in self.data:
            affine_transformation(self.data[key], self.x_hat)

    def principal_point_offest(self, coords):
        coords[0] = (coords[0] - self.camera.ppo[0])
        coords[1] = (coords[1] - self.camera.ppo[1])

    def radial_lens_correction(self, coords):
        x, y = coords

        r = math.sqrt(x**2 + y**2)
        # balanced
        dr = ((self.camera.k0 * r) + (self.camera.k1 * r**3) + (self.camera.k2 * r**5) + (self.camera.k3 * r**7))
        x_rad = -x * (dr / r)
        y_rad = -y * (dr / r)

        self.rad_table.append([r, dr, x_rad, y_rad])

        return x_rad, y_rad

    def decentering_lens_correction(self, coords):
        x, y = coords

        r = math.sqrt(x**2 + y**2)
        x_dec = -(self.camera.p1 * (r**2 + 2 * x**2)) + (2 * self.camera.p2 * x * y)
        y_dec = -(self.camera.p2 * (r**2 + 2 * y**2)) + (2 * self.camera.p1 * x * y)

        self.dec_table.append([r, x_dec, y_dec])

        return x_dec, y_dec

    # a = [], h = [], H = int
    def atmospheric_refraction_correction(self, coords, K):
        x, y = coords
        
        r = math.sqrt(x**2 + y**2)
        dr = K * r * (1 + (r**2/self.camera.f**2))
        x_atm = -x * K * (1 + (r**2/self.camera.f**2))
        y_atm = -y * K * (1 + (r**2/self.camera.f**2))

        self.atm_table.append([r, dr, x_atm, y_atm])

        return x_atm, y_atm

    def set_height(self, h, H):
        self.h = h
        self.H = H

    def get_K(self):
        h = self.h * 1E-3 # km
        H = self.H * 1E-3 # km
        K = ((2410 * H)/(H**2 - 6*H + 250)) - ((2410*h)/(h**2 - 6*h + 250)) * (h/H)
        K = K * 1E-6
        return K

    def apply_corrections(self):
        self.rad_table = []
        self.dec_table = []
        self.atm_table = []

        self.transformed = []
        self.pp_corrected = []
        self.rad_corrected = []
        self.dec_corrected = []
        self.atm_corrected = []

        self.measurements = []

        self.final = []
        self.fiducial = []
        self.tie = []
        self.control = []
        self.planar = []

        K = self.get_K()

        for key in self.data:
            for coords in self.data[key]:

                transformed_x, transformed_y = coords
                self.principal_point_offest(coords)
                pp_x, pp_y = coords

                x_rad, y_rad = self.radial_lens_correction(coords)
                x_dec, y_dec = self.decentering_lens_correction(coords)
                if (key != "fiducials"):
                    x_atm, y_atm = self.atmospheric_refraction_correction(coords, K)
                else:
                    x_atm, y_atm = 0, 0

                rad_x, rad_y = pp_x + x_rad, pp_y + y_rad
                dec_x, dec_y = rad_x + x_dec, rad_y + y_dec
                atm_x, atm_y = dec_x + x_atm, dec_y + y_atm

                test_x = pp_x + x_rad + x_dec + x_atm
                test_y = pp_y + y_rad + y_dec + y_atm

                if (test_x != atm_x or test_y != atm_y):
                    print("PROBLEM HERE: adding corrections\n")

                #print(f"TEST: {test_x - pp_x, test_y - pp_y}")

                self.measurements.append([transformed_x, transformed_y, pp_x, pp_y, rad_x, rad_y, dec_x, dec_y, atm_x, atm_y])

                # split the different points into different lists
                final_coords = [atm_x, atm_y]
                self.final.append(final_coords)
                if key == "fiducials":
                    self.fiducial.append(final_coords)
                elif key == "control":
                    self.control.append(final_coords)
                elif key == "tie":
                    self.tie.append(final_coords)
                elif key == "planar":
                    self.planar.append(final_coords)

    def correct(self, data):
        affine_transformation(data, self.x_hat)
        
        K = self.get_K()

        for coords in data:
            self.principal_point_offest(coords)
            self.radial_lens_correction(coords)
            self.decentering_lens_correction(coords)
            self.atmospheric_refraction_correction(coords, K)

doc = Document()

def dist(a, b):
    return math.sqrt((a[0] - b[0])**2 + (a[1] - b[1])**2)

def calculate_pixel_spacing(images):
    cali_dists = [
        299.816,
        299.825,
        224.005,
        224.006,
    ] # mm

    # calculate mm / px aka pixel spacing
    print("distances:")
    pixel_spacing = 0
    for i in range(4):
        cali_dist = cali_dists[i]
        it = i * 2
        pix_dist = dist(images[0].data['fiducials'][it + 1], images[0].data['fiducials'][it])
        print(f"measured: {pix_dist}")
        pixel_spacing_estimate = cali_dist / pix_dist
        print(f"spacing: {pixel_spacing_estimate}")
        pixel_spacing += pixel_spacing_estimate
    pixel_spacing = pixel_spacing / 4 # mm / px
    print(f"pixel_spacing: {pixel_spacing}")
    return pixel_spacing

def test_affine_transformation():
    test_data = [
        [-113.767,	-107.400],
        [-43.717,	-108.204],
        [36.361,	-109.132],
        [106.408,	-109.923],
        [107.189,	-39.874],
        [37.137,	-39.070],
        [-42.919,	-38.158],
        [-102.968,	-37.446],
        [-112.052,	42.714],
        [-42.005,	41.903],
        [38.051,	40.985],
        [108.089,	40.189],
        [108.884,	110.221],
        [38.846,	111.029],
        [-41.208,	111.961],
        [-111.249,	112.759],
    ]

    test_data_reseaux = [
        [-110,	-110],
        [-40,	-110],
        [40,	-110],
        [110,	-110],
        [110,	-40],
        [40,	-40],
        [-40,	-40],
        [-100,	-40],
        [-110,	40],
        [-40,	40],
        [40,	40],
        [110,	40],
        [110,	110],
        [40,	110],
        [-40,	110],
        [-110,	110],
    ]
    print("\nlecture data:")
    x_hat_test = get_affine_transformation_parameters(test_data, test_data_reseaux, "test")
    affine_transformation(data, x_hat_test)

def get_height():
    avg_h = 0
    for coords in object_coords:
        x, y, z = coords
        avg_h = avg_h + z
    avg_h = avg_h / len(object_coords)

    flying_height = 751.4637599031875 # flying height from Lab 1
    H = flying_height + avg_h
    return avg_h, H

def main():
    og_27 = {
        "fiducials": load_data(data, 0, 8),
        "control": load_data(data, 16, 24),
        "tie": load_data(data, 32, 38),
    }

    og_28 = {
        "fiducials": load_data(data, 8, 16),
        "control": load_data(data, 24, 32),
        "tie": load_data(data, 38, 44),
    }

    camera = Camera()

    images = [
        Image(camera, "image_27", 20448, 20480, og_27),
        Image(camera, "image_28", 20462, 20494, og_28)
    ]

    # part b

    images[0].pool()
    images[1].pool()

    # part c

    # fiducial marks from certificate
    images[0].get_affine_transformation_parameters()
    images[1].get_affine_transformation_parameters()

    images[0].apply_transformation()
    images[1].apply_transformation()

    # part d
    test_affine_transformation()

    # part e
    print("\n\npart e")

    avg_h, H = get_height()
    images[0].set_height(avg_h, H)
    images[1].set_height(avg_h, H)

    images[0].apply_corrections()
    images[1].apply_corrections()

    array_to_word_table(concat_2d(images[0].rad_table, images[1].rad_table), f"radial_lens", float, decimals=4)
    array_to_word_table(concat_2d(images[0].dec_table, images[1].dec_table), f"decentering_lens", float, decimals=4)
    array_to_word_table(concat_2d(images[0].atm_table, images[1].atm_table), f"atmospheric_refraction", float, decimals=4)

    final_table = []
    for i in range(len(images[0].measurements)):
        new_row = []
        for j in range(len(images[0].measurements[0])):
            if j % 2 == 0:
                new_row.append(images[0].measurements[i][j])
                new_row.append(images[0].measurements[i][j + 1])
                new_row.append(images[1].measurements[i][j])
                new_row.append(images[1].measurements[i][j + 1])

        final_table.append(new_row)

    print("\nimage 27")
    for v in images[0].measurements:
        print(f"{v[8]}, {v[9]}")
    print("\nimage 28")
    for v in images[1].measurements:
        print(f"{v[8]}, {v[9]}")

    array_to_word_table(final_table, f"measurements", float, decimals=3)

    print(np.array(images[0].final))

    return images


if __name__ == "__main__":
    main()
    doc.save("output.docx")