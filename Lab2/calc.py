import pandas as pd
import math
import numpy as np
import matplotlib.pyplot as plt
import copy

from docx import Document

'''
def dist_3D(a, b):
    return math.sqrt((a[0] - b[0])**2 + (a[1] - b[1])**2 + (a[2] - b[2])**2)

longest_dist = 0
for a in object_coords:
    for b in object_coords:
        if a is b:
            continue
        dist = dist_3D(a, b)
        if dist > longest_dist:
            longest_dist = dist
            print(f"A: {a}, B: {b}")
'''

def load_data(data, low, high):
    load = []
    for i in range(low, high):
        same = []
        for j, value in enumerate(data[i]):
            if j % 2 == 0:
                coords = [data[i][j], data[i][j+1]]
                same.append(coords)
        load.append(same)
    return load

def right_hand(a, image_dim):
    for coords in a:
        x = coords[0] - (image_dim[0]/2)
        y = (image_dim[1]/2) - coords[1]

        coords[0] = x
        coords[1] = y


def mean(a):
    sum = [0, 0]
    for v in a:
        if not math.isnan(v[0]):
            sum[0] = sum[0] + v[0]
        if not math.isnan(v[1]):
            sum[1] = sum[1] + v[1]
    sum[0] = sum[0] / len(a)
    sum[1] = sum[1] / len(a)
    return sum

def std(a):
    m = mean(a)
    sum = [0, 0]
    for v in a:
        if not math.isnan(v[0]):
            sum[0] = sum[0] + (v[0] - m[0])**2
        if not math.isnan(v[1]):
            sum[1] = sum[1] + (v[1] - m[1])**2
    sum[0] = math.sqrt(sum[0]/(len(a) - 1))
    sum[1] = math.sqrt(sum[1]/(len(a) - 1))
    return sum

def array_3d_to_2d(a):
    output = []

    for v1 in a:
        e = []
        for v2 in v1:
            for v3 in v2:
                e.append(v3)
        output.append(e)

    return output

def concat_2d(a, b):
    return [row1 + row2 for row1, row2 in zip(a, b)]

# print a 2d array to a word table
def array_to_word_table(array, name, type, decimals=2):
    doc.add_paragraph(name)
    table = doc.add_table(rows=len(array), cols=len(array[0]))

    for i, row in enumerate(array):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(type(round(cell, decimals)))

def design_matrix(fiducial):
    n = len(fiducial)
    A = np.zeros((2 * n, 6))

    for i in range(n):
        x, y = fiducial[i]
        A[2*i, 0] = x
        A[2*i, 1] = y
        A[2*i, 2] = 1
        A[2*i + 1, 3] = x
        A[2*i + 1, 4] = y
        A[2*i + 1, 5] = 1

    return A

def residual_plot(og, residuals):
    print("\n\nPlot")

    og = np.array(og)
    r = np.array(r)

    plt.figure(figsize=(8, 6))

    for i in range(len(og)):
        plt.arrow(og[i, 0], og[i, 1], r[i, 0], r[i, 1], head_width=2.5, head_length=5.5, color='red')

    plt.xlabel("x (mm)")
    plt.ylabel("y (mm)")
    #plt.axhline(y=0, color='black', linewidth=0.7, linestyle='--')
    plt.legend()
    plt.show()

def RMS(a):
    sum = 0
    for v in a:
        sum = sum + v**2
    
    return math.sqrt(sum/len(a))

def affine_transformation(data, x_hat):
    for coords in data:
        x = (x_hat[0] * coords[0] + x_hat[1] * coords[1]) + x_hat[2]
        y = (x_hat[3] * coords[0] + x_hat[4] * coords[1]) + x_hat[5]

        coords[0] = x
        coords[1] = y

# data is the fiducials in a [[]]

def get_affine_transformation_parameters(data, t_data, tag):
    print(f"\n{tag}:")
    print(data)

    l = []
    for coords in t_data:
        l.append(coords[0])
        l.append(coords[1])

    l = np.array(l)
    A = design_matrix(data)

    x_hat = np.linalg.inv(A.T @ A) @ A.T @ l
    print(f"x_hat: {x_hat}")
    # a b x c d y

    # displaying residuals in stats and graphs

    residuals = A @ x_hat - l
    print(f"residuals: {residuals}")

    # convert 1D-array into 2D-array
    r = []
    r_3d = []
    for i in range(len(residuals)):
        if i % 2 == 0:
            coord = [residuals[i], residuals[i + 1]]
            r.append(coord)
            r_3d.append([coord])

    #residual_plot(t_data, residuals)

    r = np.array(r)
    print(f"RMS x: {RMS(r[:, 0])}, y: {RMS(r[:, 1])}")

    array_to_word_table(r, f"residual_{tag}", float, decimals=6)

    return x_hat

class Camera:
    def __init__(self):
        self.ppo = [-0.006, 0.006] # mm

        self.k0 = 0.8878e-4
        self.k1 = -0.1528e-7
        self.k2 = 0.5256e-12
        self.k3 = 0

        self.p1 = 0.1346e-6
        self.p2 = 0.1224e-7

        self.f = 153.358 # focal length (mm)

        self.fiducial_marks = [
            [-105.997, -105.995],
            [106.004, 106.008],
            [-106.000, 106.009],
            [106.012, -105.995],
            [-112.000, 0.007],
            [112.006, 0.007],
            [0.005, 112.007],
            [0.002, -111.998]
        ]

class Image:
    def __init__(self, camera, tag, dim_x, dim_y, data):
        self.camera = camera
        self.tag = tag
        self.dim = [dim_x, dim_y]
        self.all_data = data

    def pool(self):
        self.data = {key: 0 for key in self.all_data.keys()}

        # applying right-handed transformation
        for key in self.all_data:
            for v in self.all_data[key]:
                right_hand(v, self.dim)

            array_to_word_table(array_3d_to_2d(self.all_data[key]), f"{key}_{self.tag}", int)
            print(f"{key}_{self.tag}\n{self.all_data[key]}")

        # calculating mean and std
        for i, key in enumerate(self.all_data):
            a = self.all_data[key]
            values = []

            table = []

            print(f"{key}:")
            for v in a:
                m = mean(v)
                s = std(v)
                print(f"mean: {m}, std: {s}")
                values.append(m)
                table.append([m[0], m[1], s[0], s[1]])

            self.data[key] = values

            array_to_word_table(table, f"mean__std_{key}_{self.tag}", float)

    def get_affine_transformation_parameters(self):
        self.x_hat = get_affine_transformation_parameters(self.data['fiducials'], self.camera.fiducial_marks, self.tag)

    def apply_transformation(self):
        for key in self.data:
            affine_transformation(self.data[key], self.x_hat)

    def principal_point_offest(self, coords):
        coords[0] = (coords[0] - self.camera.ppo[0])
        coords[1] = (coords[1] - self.camera.ppo[1])

    def radial_lens_correction(self, coords):
        x, y = coords

        r = math.sqrt(x**2 + y**2)
        # balanced
        dr = ((self.camera.k0 * r) + (self.camera.k1 * r**3) + (self.camera.k2 * r**5) + (self.camera.k3 * r**7))
        x_rad = -x * (dr / r)
        y_rad = -y * (dr / r)

        self.rad_table.append([r, dr, x_rad, y_rad])

        return x_rad, y_rad

    def decentering_lens_correction(self, coords):
        x, y = coords

        r = math.sqrt(x**2 + y**2)
        x_dec = -(self.camera.p1 * (r**2 + 2 * x**2)) + (2 * self.camera.p2 * x * y)
        y_dec = -(self.camera.p2 * (r**2 + 2 * y**2)) + (2 * self.camera.p1 * x * y)

        self.dec_table.append([r, x_dec, y_dec])

        return x_dec, y_dec

    # a = [], h = [], H = int
    def atmospheric_refraction_correction(self, coords, K):
        x, y = coords
        
        r = math.sqrt(x**2 + y**2)
        dr = K * r * (1 + (r**2/self.camera.f**2))
        x_atm = -x * K * (1 + (r**2/self.camera.f**2))
        y_atm = -y * K * (1 + (r**2/self.camera.f**2))

        self.atm_table.append([r, dr, x_atm, y_atm])

        return x_atm, y_atm

    def apply_corrections(self, h, H):
        self.rad_table = []
        self.dec_table = []
        self.atm_table = []

        self.transformed = []
        self.pp_corrected = []
        self.rad_corrected = []
        self.dec_corrected = []
        self.atm_corrected = []

        self.measurements = []

        h = h * 1E-3 # km
        H = H * 1E-3 # km
        K = ((2410 * H)/(H**2 - 6*H + 250)) - ((2410*h)/(h**2 - 6*h + 250)) * (h/H)
        K = K * 1E-6
        print(f"K: {K}")

        for key in self.data:
            for coords in self.data[key]:

                transformed_x, transformed_y = coords
                self.principal_point_offest(coords)
                pp_x, pp_y = coords

                x_rad, y_rad = self.radial_lens_correction(coords)
                x_dec, y_dec = self.decentering_lens_correction(coords)
                if (key != "fiducials"):
                    x_atm, y_atm = self.atmospheric_refraction_correction(coords, K)
                else:
                    x_atm, y_atm = 0, 0

                rad_x, rad_y = pp_x + x_rad, pp_y + y_rad
                dec_x, dec_y = rad_x + x_dec, rad_y + y_dec
                atm_x, atm_y = dec_x + x_atm, dec_y + y_atm

                test_x = pp_x + x_rad + x_dec + x_atm
                test_y = pp_y + y_rad + y_dec + y_atm

                if (test_x != atm_x or test_y != atm_y):
                    print("PROBLEM HERE: adding corrections\n")

                print(f"TEST: {test_x - pp_x, test_y - pp_y}")

                self.measurements.append([transformed_x, transformed_y, pp_x, pp_y, rad_x, rad_y, dec_x, dec_y, atm_x, atm_y])

doc = Document()
df = pd.read_excel('data.xlsx')

data = []
for index, row in df.iterrows():
    data.append(row.to_list())

og_27 = {
    "fiducials": load_data(data, 0, 8),
    "control": load_data(data, 16, 24),
    "tie": load_data(data, 32, 38),
}

og_28 = {
    "fiducials": load_data(data, 8, 16),
    "control": load_data(data, 24, 32),
    "tie": load_data(data, 38, 44),
}

camera = Camera()

images = [
    Image(camera, "image_27", 20448, 20480, og_27),
    Image(camera, "image_28", 20462, 20494, og_28)
]


# part b

images[0].pool()
images[1].pool()

# get a value for pixel size

def dist(a, b):
    return math.sqrt((a[0] - b[0])**2 + (a[1] - b[1])**2)

cali_dists = [
    299.816,
    299.825,
    224.005,
    224.006,
] # mm

# calculate mm / px aka pixel spacing
print("distances:")
pixel_spacing = 0
for i in range(4):
    cali_dist = cali_dists[i]
    it = i * 2
    pix_dist = dist(images[0].data['fiducials'][it + 1], images[0].data['fiducials'][it])
    print(f"measured: {pix_dist}")
    pixel_spacing_estimate = cali_dist / pix_dist
    print(f"spacing: {pixel_spacing_estimate}")
    pixel_spacing += pixel_spacing_estimate
pixel_spacing = pixel_spacing / 4 # mm / px
print(f"pixel_spacking: {pixel_spacing}")

# part c

# fiducial marks from certificate
images[0].get_affine_transformation_parameters()
images[1].get_affine_transformation_parameters()

images[0].apply_transformation()
images[1].apply_transformation()

# part d

test_data = [
    [-113.767,	-107.400],
    [-43.717,	-108.204],
    [36.361,	-109.132],
    [106.408,	-109.923],
    [107.189,	-39.874],
    [37.137,	-39.070],
    [-42.919,	-38.158],
    [-102.968,	-37.446],
    [-112.052,	42.714],
    [-42.005,	41.903],
    [38.051,	40.985],
    [108.089,	40.189],
    [108.884,	110.221],
    [38.846,	111.029],
    [-41.208,	111.961],
    [-111.249,	112.759],
]

test_data_reseaux = [
    [-110,	-110],
    [-40,	-110],
    [40,	-110],
    [110,	-110],
    [110,	-40],
    [40,	-40],
    [-40,	-40],
    [-100,	-40],
    [-110,	40],
    [-40,	40],
    [40,	40],
    [110,	40],
    [110,	110],
    [40,	110],
    [-40,	110],
    [-110,	110],
]
print("\nlecture data:")
x_hat_test = get_affine_transformation_parameters(test_data, test_data_reseaux, "test")
affine_transformation(data, x_hat_test)

# part e

print("\n\npart e")

object_coords = [
    [-399.28, -679.72, 1090.96],
    [109.70, -642.35, 1086.43],
    [475.55, -538.18, 1090.50],
    [517.62, -194.43, 1090.65],
    [-466.39, -542.31, 1091.55],
    [42.73, -412.19, 1090.82],
    [321.09, -667.45, 1083.49],
    [527.78, -375.72, 1092.00]
]

avg_h = 0
for coords in object_coords:
    x, y, z = coords
    avg_h = avg_h + z
avg_h = avg_h / len(object_coords)

flying_height = 751.4637599031875 # flying height from Lab 1
H = flying_height + avg_h

print(avg_h)
print(H)

images[0].apply_corrections(avg_h, H)
images[1].apply_corrections(avg_h, H)

array_to_word_table(concat_2d(images[0].rad_table, images[1].rad_table), f"radial_lens", float, decimals=4)
array_to_word_table(concat_2d(images[0].dec_table, images[1].dec_table), f"decentering_lens", float, decimals=4)
array_to_word_table(concat_2d(images[0].atm_table, images[1].atm_table), f"atmospheric_refraction", float, decimals=4)

final_table = []
for i in range(len(images[0].measurements)):
    new_row = []
    for j in range(len(images[0].measurements[0])):
        if j % 2 == 0:
            new_row.append(images[0].measurements[i][j])
            new_row.append(images[0].measurements[i][j + 1])
            new_row.append(images[1].measurements[i][j])
            new_row.append(images[1].measurements[i][j + 1])

    final_table.append(new_row)

print("\nimage 27")
for v in images[0].measurements:
    print(f"{v[8]}, {v[9]}")
print("\nimage 28")
for v in images[1].measurements:
    print(f"{v[8]}, {v[9]}")

array_to_word_table(final_table, f"measurements", float, decimals=3)

doc.save("output.docx")

